#!/usr/bin/env python3
"""生成玫瑰SaaS产品白皮书 Word 文档（加拿大市场·加币定价）"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── 全局字体 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    hs.font.color.rgb = RGBColor(0x1a, 0x56, 0x76)

# ── 封面 ──
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Rose SaaS White Paper')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x56, 0x76)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('All-in-One Digital Solution for Canadian Restaurants')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('System Build & Deployment  ·  One-Time Fee')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0x76)
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Version 1.0 ｜ {datetime.date.today().strftime("%B %Y")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ── Table of Contents ──
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Overview',
    '2. Product Ecosystem',
    '3. Feature Comparison',
    '4. Canadian Market Research & Pricing Strategy',
    '5. Service Pricing (CAD)',
    '6. Monthly Operating Costs (Merchant Responsibility)',
    '7. Technology Architecture',
    '8. Typical Use Cases',
    '9. Competitive Advantages',
    '10. FAQ',
    '11. Contact Us',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ── 1. Overview ──
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'Rose SaaS is a one-stop digital platform purpose-built for small to medium-sized '
    'restaurants across Canada. Powered by Cloudflare\'s global edge network, it delivers '
    'restaurant website building, online ordering management, AI-powered customer service, '
    'knowledge base Q&A (RAG), and multi-channel payment integration.'
)
doc.add_paragraph(
    'The system uses a four-tier product architecture (A/B/C/D), allowing restaurants to '
    'start at any level and upgrade seamlessly as their business grows — from basic '
    'digital presence to full-stack intelligent operations.'
)
doc.add_paragraph(
    'Our mission: "Give every small Canadian restaurant enterprise-grade digital capabilities '
    'at an affordable, one-time price — with zero technical skills required."'
)
doc.add_paragraph(
    'Service model: We charge a one-time system build and deployment fee. '
    'Monthly cloud operating costs are paid directly by the merchant — '
    'most of which are covered by Cloudflare\'s generous free tier and DeepSeek\'s low-cost API. '
    'No hidden fees, no annual subscriptions, no commission on orders.'
)

# ── 2. Product Ecosystem ──
doc.add_heading('2. Product Ecosystem', level=1)

doc.add_heading('Product A — Essential "Cloud Restaurant"', level=2)
doc.add_paragraph(
    'A professional restaurant website with an online ordering management system and '
    'staffed customer service (no AI). No payment processing. Ideal for restaurants '
    'that want a branded online presence and phone/dine-in ordering.'
)
for item in [
    'Responsive restaurant website (Classic / Modern templates)',
    'Menu display and management dashboard',
    'Online reservations / phone ordering',
    'Staff-managed customer service (no AI)',
    'Order management system (receive, prepare, serve)',
    'Dedicated merchant subdomain (.pages.dev)',
    'Cloudflare global CDN acceleration',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Product B — Intelligent "AI Customer Service"', level=2)
doc.add_paragraph(
    'Everything in Product A, plus LLM-powered AI customer service (DeepSeek) with '
    'RAG (Retrieval-Augmented Generation) using your own documents — menus, policies, FAQs. '
    '7×24 automatic responses. Perfect for busy restaurants looking to reduce staffing costs '
    'and improve customer response times.'
)
items_b = ['All Product A features'] + [
    'AI customer service (DeepSeek-R1 / V3)',
    'Local document knowledge base (PDF/Word/TXT)',
    'RAG-powered real-time answers from your documents',
    'Google Drive knowledge base auto-sync',
    'Vector search (Vectorize)',
    'Conversation history & analytics',
    '7×24 automated responses',
]
for item in items_b:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Product C — Full Suite "Smart Pay"', level=2)
doc.add_paragraph(
    'Everything in Product B, plus integration with Stripe (and Square as alternative). '
    'AI assistant guides customers through ordering and checkout within the chat, '
    'creating a seamless journey from inquiry to payment. Ideal for restaurants ready '
    'to accept online payments and reduce friction.'
)
items_c = ['All Product B features'] + [
    'Stripe payment integration (2.9% + $0.30 CAD)',
    'Square payment integration (2.6% + $0.10 CAD)',
    'AI-guided order taking and payment',
    'Automatic order total calculation & payment link generation',
    'Real-time payment status callbacks & order sync',
    'Multi-currency support (CAD / USD)',
]
for item in items_c:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Product D — Custom "Enterprise Edition"', level=2)
doc.add_paragraph(
    'Fully customized development for enterprise clients. Includes all A/B/C features, '
    'plus custom workflows, multi-location management, supply chain integration, '
    'daily data maintenance, and dedicated support. Separately quoted based on requirements.'
)
for item in [
    'All A + B + C features',
    'Custom development (workflows, UI, integrations)',
    'Daily data maintenance & monitoring',
    'System health monitoring & alerts',
    'Dedicated technical account manager',
    'Priority response & bug fixes',
    'Quarterly system health reports',
]:
    doc.add_paragraph(item, style='List Bullet')

# ── Optional Add-ons ──
doc.add_heading('Optional Add-Ons', level=2)
for item in [
    'AI Voice Agent (Twilio phone integration) — one-time setup $980 CAD',
    'Custom domain setup (replace default subdomain) — one-time $380 CAD',
    'Premium template customization (brand VI design) — from $2,980 CAD',
    'Annual maintenance & support (security updates, backups, tech support) — '
    '$800 CAD/yr (A/B) or $1,500 CAD/yr (C)',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ── 3. Feature Comparison ──
doc.add_heading('3. Feature Comparison', level=1)

table = doc.add_table(rows=1, cols=5)
table.style = 'Medium Shading 1 Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Feature', 'A Essential', 'B Intelligent', 'C Full Suite', 'D Enterprise']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True

rows_data = [
    ['Responsive Website', '✔', '✔', '✔', '✔'],
    ['Menu Management', '✔', '✔', '✔', '✔'],
    ['Online Reservations', '✔', '✔', '✔', '✔'],
    ['Staff Customer Service', '✔', '✔', '✔', '✔'],
    ['Order Management', '✔', '✔', '✔', '✔'],
    ['AI Customer Service', '—', '✔', '✔', '✔'],
    ['Document RAG Knowledge Base', '—', '✔', '✔', '✔'],
    ['Google Drive Sync', '—', '✔', '✔', '✔'],
    ['DeepSeek LLM', '—', '✔', '✔', '✔'],
    ['7×24 Auto Responses', '—', '✔', '✔', '✔'],
    ['Stripe Payment', '—', '—', '✔', '✔'],
    ['Square Payment', '—', '—', '✔', '✔'],
    ['AI-Guided Order & Pay', '—', '—', '✔', '✔'],
    ['Custom Development', '—', '—', '—', '✔'],
    ['Daily Data Maintenance', '—', '—', '—', '✔'],
    ['AI Voice Agent (optional)', 'Optional', 'Optional', 'Optional', 'Optional'],
    ['Custom Domain (optional)', 'Optional', 'Optional', 'Optional', 'Optional'],
]
for row_data in rows_data:
    row = table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        for paragraph in row.cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── 4. Canadian Market Research & Pricing Strategy ──
doc.add_heading('4. Canadian Market Research & Pricing Strategy', level=1)

doc.add_heading('4.1 Canadian Market Price Benchmarks (2026)', level=2)

mkt = doc.add_table(rows=1, cols=4)
mkt.style = 'Medium Shading 1 Accent 1'
mkt.alignment = WD_TABLE_ALIGNMENT.CENTER

mkt_headers = ['Service Category', 'Representative Providers', 'Pricing Model', 'Price Range (CAD)']
for i, h in enumerate(mkt_headers):
    cell = mkt.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True

mkt_rows = [
    ['Restaurant Website (basic)', 'EB Media, SOS Tech, freelancers', 'One-time build', '$1,250 – $4,000'],
    ['Restaurant Website (standard)', 'Setsail Marketing, agencies', 'One-time build', '$4,000 – $10,000'],
    ['Restaurant Website (e-commerce)', 'Specialized agencies', 'One-time build', '$10,000 – $25,000+'],
    ['Online Ordering System', 'MenuDirect, DineTech, InstaMenu', 'Monthly SaaS + setup fee', '$15 – $200/mo + $99–$580 setup'],
    ['AI Chatbot (SaaS)', 'BotHero, ChatDirect.ca, BabaAI', 'Monthly subscription', '$29 – $299/mo'],
    ['AI Chatbot + Voice', 'ThinkBizAI, BabaAI', 'Monthly subscription', '$99 – $599/mo'],
    ['Payment Processing', 'Stripe, Square, Moneris', 'Per-transaction', '1.5% – 3% + $0.10–$0.30'],
    ['Custom System Dev', 'Agencies, dev shops', 'Project-based', '$10,000 – $100,000+'],
]
for row_data in mkt_rows:
    row = mkt.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        for paragraph in row.cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('4.2 Pricing Strategy', level=2)
doc.add_paragraph(
    'We use a one-time system build & deployment fee model — not annual subscriptions. '
    'Rationale:'
)
for item in [
    'Monthly operating costs are minimal ($0–3 CAD/mo base, plus payment processing fees), '
    'so there is no need to amortize infrastructure costs through recurring fees.',
    'Canadian small business owners overwhelmingly prefer one-time payments over ongoing commitments.',
    'Market AI chatbots charge $29–$299/mo; our one-time fee matches 6–12 months of their SaaS fees '
    'but gives permanent ownership of the built system.',
    'Scope is well-defined (template → configure → deploy), making project-based pricing natural.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'Our pricing is positioned aggressively below Canadian market averages:\n'
    '• A simple website + ordering system typically costs $2,000–$4,000 in Canada.\n'
    '• AI chatbot setup + knowledge base integration typically $2,000–$5,000 extra.\n'
    '• Payment integration typically adds $1,000–$3,000.\n'
    '• Full custom systems start at $10,000+.\n\n'
    'We deliver equivalent or superior functionality at 30–50% below market rates '
    'by leveraging Cloudflare\'s free tier, DeepSeek\'s ultra-low API pricing, '
    'and our pre-built, battle-tested platform.'
)

doc.add_page_break()

# ── 5. Service Pricing (CAD) ──
doc.add_heading('5. Service Pricing (CAD)', level=1)
doc.add_paragraph(
    'All prices are in Canadian dollars (CAD). These are one-time system build and deployment fees. '
    'Monthly cloud operating costs are paid by the merchant (see Section 6).'
)

pt = doc.add_table(rows=1, cols=5)
pt.style = 'Medium Shading 1 Accent 1'
pt.alignment = WD_TABLE_ALIGNMENT.CENTER

p_headers = ['Product', 'Price (CAD)', 'Delivery Time', 'What\'s Included', 'Best For']
for i, h in enumerate(p_headers):
    cell = pt.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True

p_rows = [
    ['A — Essential\n"Cloud Restaurant"',
     '$1,980',
     '1–3 business days',
     'Template selection, menu setup, '
     'page configuration, subdomain deployment, basic training',
     'Independent restaurants, cafes, '
     'first-time digitalization'],
    ['B — Intelligent\n"AI Customer Service"',
     '$4,980',
     '3–5 business days',
     'All A + AI chat setup, knowledge base '
     'creation, document upload & training, conversation testing',
     'High-volume restaurants needing '
     'to reduce labour costs'],
    ['C — Full Suite\n"Smart Pay"',
     '$9,800',
     '5–7 business days',
     'All B + Stripe/Square integration, '
     'payment flow configuration, order-to-payment pipeline testing',
     'Restaurants ready for online '
     'payments and full automation'],
    ['D — Enterprise\n"Custom Edition"',
     '$16,800+',
     'Negotiable',
     'All A+B+C + custom development + '
     'daily maintenance + dedicated support',
     'Chains, multi-location brands, '
     'enterprises with custom needs'],
]
for row_data in p_rows:
    row = pt.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        for paragraph in row.cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph(
    'Note: Prices exclude optional add-ons (AI Voice Agent $980, Custom Domain $380, '
    'Premium Template $2,980+, Annual Maintenance $800–$1,500/yr). '
    'All prices are subject to applicable taxes (GST/HST/QST).'
)

doc.add_page_break()

# ── 6. Monthly Operating Costs ──
doc.add_heading('6. Monthly Operating Costs (Merchant Responsibility)', level=1)
doc.add_paragraph(
    'After deployment, merchants pay ongoing operating costs directly to cloud providers. '
    'We do not handle, mark up, or profit from these costs — they are completely transparent.'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Estimated monthly costs for a typical small Canadian restaurant '
                '(200–500 orders/mo, ~200 AI conversations/day, ~5,000 page views/mo)')
run.bold = True

ct = doc.add_table(rows=1, cols=5)
ct.style = 'Medium Shading 1 Accent 1'
ct.alignment = WD_TABLE_ALIGNMENT.CENTER

c_headers = ['Cost Item', 'A Essential', 'B Intelligent', 'C Full Suite', 'Billing Basis']
for i, h in enumerate(c_headers):
    cell = ct.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True

c_rows = [
    ['Cloudflare (free tier)', '$0', '$0', '$0', '100k req/day, 5GB D1, 10GB R2'],
    ['Vectorize (vector DB)', '—', '$0', '$0', 'Free during beta'],
    ['Workers AI (embeddings)', '—', '$0', '$0', 'Within free quota'],
    ['DeepSeek API (LLM tokens)', '—', '~$0.30/mo', '~$0.50/mo', 'V3: $0.14/1M in, $0.28/1M out CAD*'],
    ['Twilio Voice (optional)', '~$2.50/mo', '~$2.50/mo', '~$2.50/mo', '$1/mo number + $0.018/min CAD*'],
    ['Stripe payment processing', '—', '—', '2.9% + $0.30/txn', 'Paid per transaction'],
    ['Square payment processing', '—', '—', '2.6% + $0.10/txn', 'Paid per transaction'],
    ['—', '—', '—', '—', '—'],
    ['Monthly base cost (excl. payment)', '$0', '~$0.30', '~$0.50', 'Cloud + AI only'],
    ['Monthly cost with 300 card txns**', '—', '—', '~$145', 'Stripe: $87 + $90'],
]
for row_data in c_rows:
    row = ct.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        for paragraph in row.cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph('* Exchange rate reference: ¥1 RMB ≈ $0.19 CAD. DeepSeek API pricing converted from RMB.')
doc.add_paragraph(
    '** 300 transactions/mo × $30 CAD avg = $9,000 CAD monthly volume. '
    'Stripe: 2.9% × $9,000 = $261 + $0.30 × 300 = $90 → $351 CAD. '
    'Square: 2.6% × $9,000 = $234 + $0.10 × 300 = $30 → $264 CAD.'
)
doc.add_paragraph(
    'Key insight: Product B\'s AI operating cost is approximately $0.30 CAD/month — '
    'dramatically lower than Canadian AI chatbot SaaS alternatives ($29–$299/mo) '
    'because we use DeepSeek API on a pay-per-token basis with no monthly minimum.'
)

doc.add_page_break()

# ── 7. Technology Architecture ──
doc.add_heading('7. Technology Architecture', level=1)
doc.add_paragraph(
    'The entire system runs on Cloudflare\'s global edge network — no servers to manage, '
    'automatic scaling, built-in global CDN, and 99.99% availability.'
)

arch = [
    ('Frontend', 'Cloudflare Pages hosts the responsive SPA, global CDN, <1s first load'),
    ('API Layer', 'Cloudflare Workers edge computing, millisecond cold start'),
    ('Database', 'Cloudflare D1 (SQLite-compatible), 5GB free, automatic backups'),
    ('Storage', 'Cloudflare R2, 10GB free, for images/documents/templates'),
    ('Vector DB', 'Cloudflare Vectorize, thousand-dimension vector search for RAG'),
    ('AI Inference', 'DeepSeek API (R1/V3), as low as $0.14 CAD per million input tokens'),
    ('Embeddings', 'Workers AI (@cf/baai/bge-small-en-v1.5), within free quota'),
    ('Voice', 'Twilio Programmable Voice, pay-per-minute (~$0.018 CAD/min)'),
    ('Payments', 'Stripe & Square direct API integration'),
]
for name, desc in arch:
    p = doc.add_paragraph()
    runner = p.add_run(f'{name}: ')
    runner.bold = True
    p.add_run(desc)

doc.add_page_break()

# ── 8. Typical Use Cases ──
doc.add_heading('8. Typical Use Cases', level=1)

scenarios = [
    ('Case 1: Family Restaurant Goes Online',
     'Maria runs a family Italian restaurant in Toronto. She wants customers to see her menu '
     'and specials online. She chooses Product A ($1,980), we deploy in 2 days. '
     'Customers scan QR codes to view the menu, phone orders increase 35%. '
     'Monthly operating cost: $0 CAD.'),
    ('Case 2: Busy Burger Chain Adds AI',
     'Mike owns 3 burger joints in Vancouver. His staff spends hours answering the same '
     'questions about hours, allergens, and delivery zones. He upgrades to Product B ($4,980). '
     'After uploading his menu and FAQ documents, the AI handles 80% of inquiries. '
     'Labour cost for phone support drops 60%. Monthly AI cost: ~$0.30 CAD.'),
    ('Case 3: Pizzeria Goes Full Digital',
     'Chen\'s pizzeria in Calgary takes phone orders but 10% have mistakes due to miscommunication. '
     'He chooses Product C ($9,800). AI takes orders in chat, confirms items, calculates totals, '
     'and generates a Stripe payment link — all without human involvement. '
     'Order errors drop to 1%, revenue grows 20%. Monthly cost: ~$145 CAD (incl Stripe fees).'),
    ('Case 4: Multi-Location Brand Customization',
     'A Montreal-based fast-casual chain with 8 locations needs a unified system with custom '
     'reporting and supply chain integration. Product D ($16,800+) — custom dashboards, '
     'cross-location analytics, and dedicated maintenance team.'),
]
for title, desc in scenarios:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ── 9. Competitive Advantages ──
doc.add_heading('9. Competitive Advantages', level=1)

advantages = [
    ('Zero Server Management',
     'Cloudflare edge network — no servers, no DevOps, automatic scaling.'),
    ('One-Time Fee, No Lock-In',
     'Pay once, own the system forever. No annual contracts, no surprise renewals. '
     'Cancel anytime — your system stays live.'),
    ('Dramatically Lower AI Costs',
     'Canadian AI chatbot competitors charge $29–$299 CAD/month. '
     'Our DeepSeek-powered AI costs ~$0.30 CAD/month for the same volume — '
     'over 99% savings because we use pay-per-token pricing, not fixed monthly tiers.'),
    ('Commission-Free Ordering',
     'Unlike DoorDash/Uber Eats/Skip (25–30% commission), we charge $0 per order. '
     'Payment processor fees (2.6–2.9%) are the only transaction cost — far lower than delivery apps.'),
    ('Seamless Upgrades',
     'Move from A → B → C → D without data migration. Your configuration is preserved.'),
    ('Data Privacy & Canadian Compliance',
     'Merchant data is stored in isolated Cloudflare sub-accounts. '
     'Compatible with PIPEDA requirements.'),
    ('Fast Deployment',
     'From order to live website: as fast as 1 business day. No technical skills needed '
     'from the merchant.'),
]
for title, desc in advantages:
    p = doc.add_paragraph()
    runner = p.add_run(f'{title}: ')
    runner.bold = True
    runner.font.color.rgb = RGBColor(0x1a, 0x56, 0x76)
    p.add_run(desc)

doc.add_page_break()

# ── 10. FAQ ──
doc.add_heading('10. FAQ', level=1)

faqs = [
    ('Q: Do I pay monthly or annually?',
     'A: Neither. Our fee is a one-time payment for building and deploying your system. '
     'After that, you only pay the cloud providers directly — most of which have generous free '
     'tiers. Optional annual maintenance ($800–$1,500/yr) covers updates, backups, and support.'),
    ('Q: What if I have no technical background?',
     'A: No problem. We handle the entire setup — you just provide your menu, photos, and '
     'basic business info. After deployment, you manage daily operations through a simple dashboard.'),
    ('Q: Can the AI give wrong answers?',
     'A: The AI responds strictly based on documents you upload (menu, policies, FAQs). '
     'It stays within the boundaries of your knowledge base. You can review conversation logs '
     'anytime from the admin dashboard.'),
    ('Q: Is payment processing secure?',
     'A: Yes. We integrate directly with Stripe and Square\'s official APIs. '
     'Payment data never touches our servers — transactions happen within the payment '
     'processor\'s PCI-compliant environment.'),
    ('Q: Can I upgrade later?',
     'A: Absolutely. From A → B → C → D, upgrade by paying the price difference. '
     'All your data (menu, orders, knowledge base) is preserved — no migration needed.'),
    ('Q: How does this compare to DoorDash/Uber Eats?',
     'A: Those platforms charge 25–30% commission per order. Our system costs $0/order — '
     'you only pay the standard credit card processing fee (2.6–2.9%). On $9,000 monthly '
     'volume, that saves you $2,250+ CAD/month vs. delivery apps.'),
    ('Q: What if I need help after deployment?',
     'A: Optional annual maintenance covers security patches, content backups, and technical '
     'support. For Product D, daily maintenance and dedicated support are included.'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    runner = p.add_run(q)
    runner.bold = True
    p = doc.add_paragraph(a)
    p.paragraph_format.space_after = Pt(10)

doc.add_page_break()

# ── 11. Contact ──
doc.add_heading('11. Contact Us', level=1)
doc.add_paragraph('Rose SaaS — Digital solutions for Canadian restaurants.')
doc.add_paragraph('Website: https://saas.roseai.ca')
doc.add_paragraph('Email: support@roseai.ca')
doc.add_paragraph('Business inquiries: business@roseai.ca')

doc.add_paragraph()
p = doc.add_paragraph('— End of Document —')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ──
output_path = '/Users/sean/coding/roseai/rose-saas/Rose_SaaS_Whitepaper.docx'
doc.save(output_path)
print(f'Whitepaper generated: {output_path}')
