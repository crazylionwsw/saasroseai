#!/usr/bin/env python3
"""Generate Product Requirements Specification (PRD) in Word format."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)

for level in range(1, 4):
    heading = doc.styles[f'Heading {level}']
    heading.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ══════════════════════════════
# COVER PAGE
# ══════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('产品需求规格说明书\nProduct Requirements Specification')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Rose SaaS — 餐饮门店在线智能服务体平台')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_info = [
    f'文档版本: V1.0',
    f'生成日期: {datetime.date.today().strftime("%Y-%m-%d")}',
    f'文档状态: 初稿',
    f'密级: 内部',
]
for line in meta_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ══════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    '1. 引言',
    '  1.1 项目背景',
    '  1.2 项目目标',
    '  1.3 适用范围',
    '  1.4 术语与缩写',
    '2. 产品概述',
    '  2.1 产品定位',
    '  2.2 目标用户',
    '  2.3 核心价值',
    '  2.4 系统架构总览',
    '3. 功能需求',
    '  3.1 中央管控平台（Admin）',
    '    3.1.1 认证与登录',
    '    3.1.2 仪表盘',
    '    3.1.3 商户管理',
    '    3.1.4 模板管理',
    '    3.1.5 网站采集（Template Scraper）',
    '    3.1.6 部署管理',
    '    3.1.7 CF 直连部署',
    '    3.1.8 部署历史',
    '    3.1.9 系统设置',
    '  3.2 商户官网系统',
    '    3.2.1 模板系统',
    '    3.2.2 官网生成',
    '    3.2.3 在线菜单展示',
    '    3.2.4 联系方式',
    '  3.3 智能客服系统（Chat）',
    '  3.4 智能电话系统（Phone）',
    '  3.5 订单系统',
    '  3.6 支付系统',
    '  3.7 Google Drive 知识库',
    '4. 非功能需求',
    '  4.1 安全性',
    '  4.2 性能',
    '  4.3 可用性',
    '  4.4 可维护性',
    '5. 数据设计',
    '  5.1 中央数据库',
    '  5.2 商户数据库',
    '6. API 设计',
    '  6.1 公开接口',
    '  6.2 商户管理接口',
    '  6.3 部署接口',
    '  6.4 模板接口',
    '7. 部署架构',
    '8. 费用估算',
    '9. 未来规划',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    for r in p.runs:
        r.font.size = Pt(10)
        if not item.startswith('  '):
            r.bold = True

doc.add_page_break()

# ══════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════
doc.add_heading('1. 引言', level=1)

doc.add_heading('1.1 项目背景', level=2)
doc.add_paragraph(
    '中小型餐饮门店普遍缺乏数字化能力：没有官方网站、没有在线下单系统、没有智能客服。'
    '现有的建站平台费用高昂，而免费方案要么功能有限，要么操作复杂。'
    '随着 AI 技术的普及，餐饮门店亟需一个低成本、易用、智能化的在线服务方案。'
)
doc.add_paragraph(
    '本项目旨在利用 Cloudflare 免费服务套件，为餐饮门店构建一套 SaaS 平台，'
    '提供官网模板、在线点餐、AI 智能客服（文本 + 语音）、知识库 RAG 等功能，'
    '所有基础设施均使用 Cloudflare 免费额度，仅语音电话环节依赖 Twilio（约 $1-2/商户/月）。'
)

doc.add_heading('1.2 项目目标', level=2)
doc.add_paragraph(
    '构建一套多租户 SaaS 平台，使餐饮门店能够：'
)
goals = [
    '快速拥有专业美观的官方网站（模板化、自动生成）',
    '实现在线菜单展示和下单功能',
    '通过 AI 智能客服（文本聊天 + 语音电话）提升顾客体验',
    '通过 Google Drive 知识库自动同步菜单、营业信息等',
    '商户零运维，所有基础设施由中央平台统一管理',
    '完全基于 Cloudflare 免费服务，仅语音通话产生少量费用',
]
for g in goals:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('1.3 适用范围', level=2)
doc.add_paragraph(
    '本 PRD 涵盖 Rose SaaS 平台 V1.0 版本的全部功能需求，包括中央管控平台、'
    '商户官网生成系统、智能客服系统、部署系统等。适用对象包括产品团队、开发团队、测试团队。'
)

doc.add_heading('1.4 术语与缩写', level=2)
add_table(doc,
    ['术语', '说明'],
    [
        ['SaaS', 'Software as a Service，软件即服务'],
        ['商户', '使用本平台的餐饮门店'],
        ['DO', 'Durable Object，Cloudflare 有状态对象'],
        ['RAG', 'Retrieval-Augmented Generation，检索增强生成'],
        ['JWT', 'JSON Web Token，用于认证的令牌'],
        ['D1', 'Cloudflare 的 SQLite 数据库服务'],
        ['R2', 'Cloudflare 的对象存储服务'],
        ['Vectorize', 'Cloudflare 的向量检索服务'],
        ['Workers AI', 'Cloudflare 的 AI 推理服务'],
        ['Pages', 'Cloudflare 的静态网站托管服务'],
        ['STT', 'Speech-to-Text，语音转文字'],
        ['TTS', 'Text-to-Speech，文字转语音'],
        ['IVR', 'Interactive Voice Response，交互式语音应答'],
        ['CF', 'Cloudflare'],
        ['API Token', 'Cloudflare API 访问令牌'],
        ['中央管控平台/Admin', '平台管理后台 UI'],
    ]
)

doc.add_page_break()

# ══════════════════════════════
# 2. PRODUCT OVERVIEW
# ══════════════════════════════
doc.add_heading('2. 产品概述', level=1)

doc.add_heading('2.1 产品定位', level=2)
doc.add_paragraph(
    'Rose SaaS 是一个面向中小餐饮门店的多租户 SaaS 平台，提供一站式数字化解决方案：'
    '从官网建站、在线点餐到 AI 客服，所有功能开箱即用。'
    '平台本身完全免费，商户仅需承担可选增值服务费用。'
)

doc.add_heading('2.2 目标用户', level=2)
add_table(doc,
    ['用户角色', '描述', '核心需求'],
    [
        ['平台管理员', 'SaaS 平台运营方', '管理商户、监控系统、配置模板、审计日志'],
        ['餐饮商户', '餐厅老板/管理者', '快速建站、在线接单、AI 客服、知识库同步'],
        ['顾客', '餐厅消费者', '浏览菜单、在线点餐、咨询客服'],
    ]
)

doc.add_heading('2.3 核心价值', level=2)
values = [
    ('零成本起步', '全部基础设施基于 Cloudflare 免费额度，仅语音电话产生微量费用'),
    ('一键部署', '填入商户 CF API Token 即可自动部署完整站点'),
    ('智能客服', 'AI 驱动的文本 + 语音客服，支持 RAG 知识库'),
    ('模板丰富', '内置经典/现代模板 + 网站自动采集生成模板'),
    ('安全可靠', 'JWT 认证、速率限制、SQL 注入防护、审计日志'),
    ('独立隔离', '每个商户独立 CF 账号，配额互不影响'),
]
for title, desc in values:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}：')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.4 系统架构总览', level=2)
doc.add_paragraph('整体架构采用中央管控 + 商户独立账号的模式：')
add_code_block(doc, '''┌─────────────────────────────────────────────────────────┐
│              中央管控平台 (Cloudflare 账号 B)              │
│    ┌──────────┐  ┌─────────────┐  ┌────────────────┐    │
│    │ Admin UI │  │ Central API │  │ D1 / R2 / AI  │    │
│    │ (Pages)  │  │ (Worker)    │  │ (平台资源)     │    │
│    └──────────┘  └──────┬──────┘  └────────────────┘    │
└─────────────────────────┼────────────────────────────────┘
                          │ HTTPS + JWT
                          ▼
┌─────────────────────────────────────────────────────────┐
│              商户 A (独立 Cloudflare 账号)                │
│  ┌──────────────────┐  ┌────────────────────────────┐   │
│  │ Storefront Pages  │  │  Merchant Worker           │   │
│  │ (静态模板网站)    │  │  /orders, /chat, /phone   │   │
│  └──────────────────┘  └───────────┬────────────────┘   │
│                                    │                    │
│  ┌──────┐  ┌────────┐  ┌───────────┐  ┌──────────┐    │
│  │ D1   │  │ R2     │  │ Vectorize │  │ Chat DO  │    │
│  │ 订单  │  │ 图片/录音│  │ 知识库向量 │  │ Phone DO │    │
│  └──────┘  └────────┘  └───────────┘  └──────────┘    │
└─────────────────────────────────────────────────────────┘''')

add_table(doc,
    ['层级', '技术', '费用'],
    [
        ['Admin UI', 'Cloudflare Pages + 原生 JS', '免费'],
        ['Central API', 'Cloudflare Workers + itty-router', '免费'],
        ['中央数据库', 'Cloudflare D1 (SQLite)', '免费(5GB)'],
        ['模板存储', 'Cloudflare R2', '免费(10GB)'],
        ['认证', 'JWT + Admin Token', '免费'],
        ['速率限制', '内存 Map + IP', '免费'],
        ['审计', 'D1 audit_logs 表', '免费'],
        ['语音电话', 'Twilio API', '~$1-2/商户/月'],
    ]
)

doc.add_page_break()

# ══════════════════════════════
# 3. FUNCTIONAL REQUIREMENTS
# ══════════════════════════════
doc.add_heading('3. 功能需求', level=1)

# ── 3.1 中央管控平台 ──
doc.add_heading('3.1 中央管控平台（Admin）', level=2)
doc.add_paragraph(
    '中央管控平台是 SaaS 系统的管理后台，部署于 Cloudflare Pages，'
    '通过 Central API Worker 与后端交互。管理员通过 Admin Token + JWT 进行双重认证。'
)

doc.add_heading('3.1.1 认证与登录', level=3)
doc.add_paragraph('功能描述：管理员输入 Admin Token 登录系统，获取 JWT Session Token。')
add_table(doc,
    ['功能点', '说明', '优先级'],
    [
        ['密码式 Token 输入', '使用 <input type="password"> 替代 prompt()', 'P0'],
        ['Session 有效期配置', 'Settings 页面配置 1/3/7/14/30 天', 'P0'],
        ['速率限制', '登录接口 5 次/分钟/IP', 'P0'],
        ['审计日志', 'LOGIN_SUCCESS / LOGIN_FAIL 记录', 'P0'],
        ['自动登出', 'JWT 过期后清除 token 并跳转登录页', 'P0'],
    ]
)

doc.add_heading('3.1.2 仪表盘', level=3)
doc.add_paragraph('功能描述：展示平台概览数据，包括商户数量、活跃商户、部署统计等。')
add_table(doc,
    ['功能点', '说明', '优先级'],
    [
        ['商户总数', '显示平台所有商户数（不含已删除）', 'P0'],
        ['活跃商户数', '状态为 active 的商户数量', 'P0'],
        ['部署次数', '总部署记录数', 'P1'],
        ['状态分布', '按状态显示商户分布', 'P1'],
    ]
)

doc.add_heading('3.1.3 商户管理', level=3)
doc.add_paragraph('功能描述：平台管理员对商户进行增删改查操作。')
add_table(doc,
    ['功能点', '说明', '优先级'],
    [
        ['商户列表', '分页展示商户，显示名称、子域名、状态、套餐、创建时间', 'P0'],
        ['搜索/筛选', '按名称/子域名搜索，按状态筛选', 'P0'],
        ['新增商户', '填写名称、电话、邮箱、套餐、模板等信息', 'P0'],
        ['编辑商户', '修改商户信息、CF 配置、模板选择', 'P0'],
        ['商户详情', '查看基本信息、Token 管理、部署历史', 'P0'],
        ['冻结/激活', '切换商户状态以控制服务可用性', 'P0'],
        ['软删除', '标记删除而非物理删除', 'P0'],
        ['CF 配置', '输入商户的 Cloudflare Account ID 和 API Token（掩码显示）', 'P0'],
        ['站点访问', '商户列表中子域名可点击直达商户站点', 'P0'],
        ['Token 重新生成', '签发新的商户 JWT Token', 'P0'],
        ['部署按钮', '从管理端触发部署', 'P1'],
    ]
)

doc.add_heading('3.1.4 模板管理', level=3)
doc.add_paragraph('功能描述：管理商户可用的网站模板。')
add_table(doc,
    ['功能点', '说明', '优先级'],
    [
        ['模板列表', '展示所有模板（内置 + 采集的），含 ID、名称、类型、状态', 'P0'],
        ['内置模板', 'classic（经典）、modern（现代）两个内置模板', 'P0'],
        ['采集模板', '通过网站采集生成的模板', 'P0'],
        ['模板状态', '启用/停用标识', 'P1'],
    ]
)

doc.add_heading('3.1.5 网站采集（Template Scraper）', level=3)
doc.add_paragraph(
    '功能描述：从指定餐厅网站 URL 自动抓取内容并生成可用的网站模板。'
)
doc.add_paragraph('采集流程：')
steps = [
    '管理员输入餐厅网站 URL',
    '系统自动抓取首页及发现的内页（最多 8 个页面）',
    '正则提取餐厅名称、电话、邮箱、地址、营业时间、简介、口号、社交媒体链接',
    '从 HTML/CSS 中提取主色调',
    '识别菜单页面，提取菜品名称和价格',
    '生成完整的 4 文件模板（index.html, menu.html, contact.html, style.css）',
    '模板文件存入 R2（templates/{id}/ + scraped/{id}/）',
    '注册模板信息到 D1 templates 表',
    '返回采集结果和模板 ID',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

add_table(doc,
    ['功能点', '说明', '优先级'],
    [
        ['输入 URL', '管理员输入餐厅官网地址', 'P0'],
        ['自动抓取', '抓取主页及内页内容', 'P0'],
        ['信息提取', '正则提取餐厅关键信息', 'P0'],
        ['模板生成', '生成含占位符 {{VAR}} 的模板文件', 'P0'],
        ['R2 存储', '模板文件持久化存储', 'P0'],
        ['D1 注册', '模板信息写入数据库', 'P0'],
        ['速率限制', '每个 IP 3 次/分钟的采集频率', 'P0'],
        ['任务列表', '查看所有已完成采集任务', 'P1'],
    ]
)

doc.add_heading('3.1.6 部署管理', level=3)
doc.add_paragraph('功能描述：记录和管理部署操作。')
add_table(doc,
    ['功能点', '说明', '优先级'],
    [
        ['创建部署', '记录版本号、状态、开始时间', 'P0'],
        ['部署历史列表', '按商户查看部署记录', 'P0'],
        ['状态更新', '支持 pending → deploying → success/failed', 'P0'],
        ['URL 记录', '记录 worker_url 和 pages_url', 'P1'],
    ]
)

doc.add_heading('3.1.7 CF 直连部署', level=3)
doc.add_paragraph(
    '功能描述：使用商户配置的 Cloudflare API Token，通过 Cloudflare Pages Direct Upload API '
    '直接将商户网站部署到商户的 Cloudflare 账号。'
)
doc.add_paragraph('部署流程：')
steps = [
    '从 D1 读取商户配置的 cf_account_id 和 cf_api_token',
    '从 R2 读取商户使用的模板文件',
    '将模板占位符（{{RESTAURANT_NAME}}, {{PHONE}} 等）替换为商户数据',
    '检查商户 CF 账号下 Pages project 是否存在，不存在则创建',
    '通过 Direct Upload API 创建 deployment（manifest + upload URLs）',
    '上传每个文件到 presigned URL',
    '在 deployments 表中记录部署结果',
    '返回部署成功的 Pages URL',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

add_table(doc,
    ['功能点', '说明', '优先级'],
    [
        ['读取 CF 配置', '从商户信息获取 Account ID 和 API Token', 'P0'],
        ['模板渲染', '替换占位符变量', 'P0'],
        ['Pages Project 创建', '自动创建 storefront-{id} 项目', 'P0'],
        ['Direct Upload', '使用 Cloudflare API 部署文件', 'P0'],
        ['部署记录', '写入 deployments 表', 'P0'],
        ['错误处理', '配置缺失时优雅报错', 'P0'],
    ]
)

doc.add_heading('3.1.8 部署历史', level=3)
doc.add_paragraph('功能描述：在商户详情页查看该商户的部署历史记录。')
doc.add_paragraph('展示内容：版本号、状态、开始/结束时间、日志查看按钮。')

doc.add_heading('3.1.9 系统设置', level=3)
doc.add_paragraph('功能描述：Session 有效期配置。')
add_table(doc,
    ['功能点', '说明', '优先级'],
    [
        ['Session 天数', '下拉选择 1/3/7/14/30 天，存入 localStorage', 'P0'],
        ['重新登录', 'JWT 过期后自动跳转登录', 'P0'],
    ]
)

doc.add_page_break()

# ── 3.2 商户官网系统 ──
doc.add_heading('3.2 商户官网系统', level=2)

doc.add_heading('3.2.1 模板系统', level=3)
doc.add_paragraph('功能描述：每个模板为独立静态文件夹，存储在 R2。模板文件使用 {{VAR}} 占位符。')
add_code_block(doc, '''templates/
├── classic/               # 内置：经典餐厅模板
│   ├── index.html         # 首页（含 {{RESTAURANT_NAME}} 等占位符）
│   ├── menu.html          # 菜单页
│   ├── contact.html       # 联系方式页
│   └── style.css          # 样式
├── modern/                # 内置：现代风格模板
│   └── ...
└── tpl-xxxx/              # 采集生成的模板
    └── ...''')

doc.add_paragraph('模板占位符规范：')
add_table(doc,
    ['占位符', '说明', '来源'],
    [
        ['{{RESTAURANT_NAME}}', '餐厅名称', '商户信息'],
        ['{{RESTAURANT_SLOGAN}}', '餐厅口号', '采集/商户'],
        ['{{RESTAURANT_DESC_SHORT}}', '简短描述', '采集/商户'],
        ['{{RESTAURANT_DESC}}', '详细描述', '采集/商户'],
        ['{{PHONE}}', '联系电话', '商户信息'],
        ['{{EMAIL}}', '联系邮箱', '商户信息'],
        ['{{ADDRESS}}', '地址', '采集/商户'],
        ['{{BUSINESS_HOURS}}', '营业时间', '采集/商户'],
        ['{{LOGO_URL}}', 'Logo 图片 URL', '待扩展'],
        ['{{COVER_URL}}', '封面图片 URL', '待扩展'],
        ['{{YEAR}}', '当前年份', '自动生成'],
    ]
)

doc.add_heading('3.2.2 官网生成', level=3)
doc.add_paragraph(
    '功能描述：根据商户信息和所选模板，渲染完整的静态官网。'
    '当前通过 Merchant Worker 运行时替换占位符，CF 部署时在 Central API 端渲染为静态文件。'
)

doc.add_heading('3.2.3 在线菜单展示', level=3)
doc.add_paragraph(
    '功能描述：在菜单页面展示餐厅的菜品分类和价格信息。'
    '采集的菜单数据渲染到模板中的菜单部分。'
)

doc.add_heading('3.2.4 联系方式', level=3)
doc.add_paragraph(
    '功能描述：联系方式页显示商户的电话、邮箱、地址、社交媒体链接。'
    '电话可点击拨打，邮箱可点击发送。'
)

doc.add_page_break()

# ── 3.3 智能客服系统 ──
doc.add_heading('3.3 智能客服系统（Chat）', level=2)
doc.add_paragraph(
    '功能描述：基于 WebSocket + Durable Object 的 AI 智能客服系统，'
    '支持向量检索增强生成（RAG），可接入 Google Drive 知识库。'
)
doc.add_paragraph('核心流程：')
steps = [
    '顾客通过 WebSocket 连接 Chat DO',
    '顾客发送消息 → Chat DO 接收',
    '向量检索：将消息转为 Embedding（BGE Small），从 Vectorize 检索 Top 5 相关片段',
    'Prompt 组装：System Prompt + 检索片段 + 对话历史 + 用户消息',
    'LLM 生成回复：使用 Workers AI Llama 3.1 8B',
    '回复推送给顾客',
    '支持转人工客服模式',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

add_table(doc,
    ['功能点', '说明', '优先级'],
    [
        ['WebSocket 连接', '顾客通过 WS 连接到 Chat DO', 'P0'],
        ['AI 自动回复', '基于 RAG + LLM 自动回复', 'P0'],
        ['RAG 知识库', '从 Vectorize 检索相关文档片段', 'P0'],
        ['对话历史', 'DO 内存中保持对话上下文', 'P0'],
        ['转人工', 'AI 无法处理时转接人工客服', 'P1'],
        ['会话摘要', '关闭会话时生成摘要', 'P2'],
    ]
)

doc.add_page_break()

# ── 3.4 智能电话系统 ──
doc.add_heading('3.4 智能电话系统（Phone）', level=2)
doc.add_paragraph(
    '功能描述：基于 Twilio Voice + Workers AI 的智能语音电话系统，'
    '支持语音识别（STT）、意图识别、RAG 回复、语音合成（TTS）。'
)
doc.add_paragraph('呼叫流程：')
steps = [
    '顾客拨打商户的 Twilio 号码',
    'Twilio 通过 Webhook 通知 Phone DO',
    'Phone DO 查询商户营业状态和知识库配置',
    'Twilio <Gather> 等待顾客语音输入',
    '顾客说话 → Twilio 发送录音 → Whisper STT 转文字',
    'Llama 3 意图识别：查询菜单/下单/营业时间/地址/转人工',
    'RAG 检索 + LLM 生成回答',
    'TTS 合成语音 → Twilio <Say> 播报',
    '继续对话或挂断',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

add_table(doc,
    ['功能点', '说明', '优先级'],
    [
        ['来电接听', 'Twilio Webhook 触发', 'P0'],
        ['语音识别', 'Whisper STT（Workers AI）', 'P0'],
        ['意图识别', 'Llama 3 分类用户意图', 'P0'],
        ['RAG 回复', '知识库检索 + LLM 生成', 'P0'],
        ['语音合成', 'Workers AI TTS', 'P0'],
        ['转人工', '转接指定号码', 'P1'],
        ['录音存储', '通话录音存入 R2', 'P2'],
        ['对话摘要', '挂断后生成摘要', 'P2'],
    ]
)

doc.add_page_break()

# ── 3.5 订单系统 ──
doc.add_heading('3.5 订单系统', level=2)
doc.add_paragraph(
    '功能描述：顾客在线浏览菜单、添加购物车、提交订单。'
)
add_table(doc,
    ['功能点', '说明', '优先级'],
    [
        ['菜单浏览', '按分类展示菜品', 'P1'],
        ['购物车', '前端 localStorage 管理', 'P1'],
        ['提交订单', 'POST /api/order/create', 'P1'],
        ['订单状态', 'pending → confirmed → preparing → delivering → completed', 'P1'],
        ['订单通知', 'WebSocket 推送商户', 'P2'],
    ]
)

# ── 3.6 支付系统 ──
doc.add_heading('3.6 支付系统', level=2)
add_table(doc,
    ['功能点', '说明', '优先级'],
    [
        ['微信支付', 'JSAPI 封装', 'P2'],
        ['支付宝', 'SDK 封装', 'P2'],
        ['货到付款', '平台记录', 'P1'],
        ['支付回调', '更新订单状态', 'P2'],
    ]
)

doc.add_page_break()

# ── 3.7 Google Drive 知识库 ──
doc.add_heading('3.7 Google Drive 知识库', level=2)
doc.add_paragraph(
    '功能描述：商户授权 Google Drive 后，系统自动同步指定文件夹中的文档，'
    '解析为文本后向量化存入 Vectorize，用于 AI 客服的 RAG 检索。'
)
doc.add_paragraph('同步流程：')
steps = [
    '商户通过 OAuth 2.0 授权 Google Drive（只读权限）',
    '商户选择要同步的文件夹',
    'Worker Cron Trigger 每 5 分钟检查变更',
    '检测到新文件/变更 → 下载并解析（PDF/DOCX/TXT/Google Docs）',
    '文本分块（512 tokens/块，150 tokens 重叠）',
    'BGE Small 生成 Embedding',
    '存入 Vectorize（按商户隔离）',
    '更新 D1 文件映射表和同步日志',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

add_table(doc,
    ['功能点', '说明', '优先级'],
    [
        ['OAuth 授权', 'Google Drive API 只读授权', 'P1'],
        ['文件夹选择', '商户选择同步目录', 'P1'],
        ['文件解析', 'PDF/DOCX/TXT/Google Docs 转文本', 'P1'],
        ['增量同步', '仅同步变更文件', 'P1'],
        ['向量化', 'BGE Small 生成 Embedding', 'P1'],
        ['向量存储', '存入 Vectorize', 'P1'],
        ['同步日志', '记录同步状态和统计', 'P2'],
    ]
)

doc.add_page_break()

# ══════════════════════════════
# 4. NON-FUNCTIONAL REQUIREMENTS
# ══════════════════════════════
doc.add_heading('4. 非功能需求', level=1)

doc.add_heading('4.1 安全性', level=2)
add_table(doc,
    ['需求', '实现方式', '优先级'],
    [
        ['口令安全', 'Timing-safe 比较（timingSafeEqual），防时序攻击', 'P0'],
        ['速率限制', '登录 5/min、敏感操作 20/min、采集 3/min、API 100/min，基于 IP', 'P0'],
        ['错误信息脱敏', '不泄露 SQL、HMAC、内部路径', 'P0'],
        ['SQL 注入防护', '参数化查询，禁止 SELECT *', 'P0'],
        ['列白名单', 'MERCHANT_SAFE_COLUMNS 只允许指定列', 'P0'],
        ['CORS 限制', '只允许已知的 Origin', 'P0'],
        ['安全响应头', 'X-Content-Type-Options, X-Frame-Options, Referrer-Policy', 'P0'],
        ['CSRF 防护', '要求 X-Requested-With: XMLHttpRequest', 'P0'],
        ['XSS 防护', '商户名称 HTML 标签过滤 + escHtml() 输出编码', 'P0'],
        ['Token 掩码', 'CF API Token 返回时只显示首尾各 4 位', 'P0'],
        ['API 认证', '除 /api/health /api/merchants/verify /api/auth/login 外均需 Bearer Token', 'P0'],
        ['审计日志', '所有 CRUD/登录/部署操作记录到 audit_logs', 'P0'],
    ]
)

doc.add_heading('4.2 性能', level=2)
add_table(doc,
    ['需求', '指标', '优先级'],
    [
        ['API 响应时间', '95% 请求 < 500ms', 'P0'],
        ['网站采集时间', '同步采集 < 10s', 'P0'],
        ['CF 部署时间', '包含上传 < 15s', 'P0'],
        ['并发请求', '支持 100 并发', 'P1'],
    ]
)

doc.add_heading('4.3 可用性', level=2)
add_table(doc,
    ['需求', '指标', '优先级'],
    [
        ['平台可用性', '99.9%（Cloudflare 保证）', 'P0'],
        ['数据持久性', 'D1 + R2 自动备份', 'P0'],
        ['部署回滚', '保留部署历史，可重新部署旧版本', 'P1'],
    ]
)

doc.add_heading('4.4 可维护性', level=2)
add_table(doc,
    ['需求', '实现方式', '优先级'],
    [
        ['CI/CD', 'GitHub Actions 自动部署（push 到 main）', 'P0'],
        ['模块化', '功能模块分离（merchants、deployments、cf-deploy）', 'P0'],
        ['测试', 'Vitest 单元测试（当前 24 个测试用例）', 'P0'],
        ['日志', '审计日志追踪所有操作', 'P0'],
    ]
)

doc.add_page_break()

# ══════════════════════════════
# 5. DATA DESIGN
# ══════════════════════════════
doc.add_heading('5. 数据设计', level=1)

doc.add_heading('5.1 中央数据库（D1）', level=2)
doc.add_paragraph('中央管控平台的 D1 数据库包含以下表：')

tables = [
    ('merchants', '商户主表', [
        ('id', 'TEXT PK', '商户唯一标识（m-{uuid8}）'),
        ('name', 'TEXT NOT NULL', '商户名称'),
        ('email', 'TEXT', '联系邮箱'),
        ('phone', 'TEXT', '联系电话'),
        ('status', 'TEXT', 'active/frozen/expired/deleted'),
        ('plan', 'TEXT', 'basic/pro/enterprise'),
        ('cf_account_email', 'TEXT', '商户 CF 账号邮箱'),
        ('cf_account_id', 'TEXT', '商户 CF 账号 ID'),
        ('cf_api_token', 'TEXT', '商户 CF API Token'),
        ('subdomain', 'TEXT UNIQUE', '子域名（shop-{id}）'),
        ('template_id', 'TEXT', '当前使用的模板 ID'),
        ('theme_color', 'TEXT', '主题色'),
        ('created_at', 'TEXT', '创建时间'),
        ('expires_at', 'TEXT', '过期时间'),
        ('notes', 'TEXT', '备注'),
    ]),
    ('merchant_tokens', '商户 JWT Token', [
        ('merchant_id', 'TEXT PK', '商户 ID'),
        ('token_hash', 'TEXT NOT NULL', 'Token 哈希'),
        ('issued_at', 'TEXT', '签发时间'),
        ('expires_at', 'TEXT', '过期时间'),
        ('last_verified_at', 'TEXT', '最后验证时间'),
    ]),
    ('deployments', '部署记录', [
        ('id', 'TEXT PK', '部署 ID（d-{uuid8}）'),
        ('merchant_id', 'TEXT NOT NULL', '商户 ID'),
        ('version', 'TEXT NOT NULL', '版本号'),
        ('template_version', 'TEXT', '模板版本'),
        ('status', 'TEXT', 'pending/deploying/success/failed'),
        ('worker_url', 'TEXT', 'Worker URL'),
        ('pages_url', 'TEXT', 'Pages URL'),
        ('cf_deployment_id', 'TEXT', 'CF 部署 ID'),
        ('started_at', 'TEXT', '开始时间'),
        ('completed_at', 'TEXT', '结束时间'),
        ('error_log', 'TEXT', '错误日志'),
        ('deployed_by', 'TEXT', '部署人'),
    ]),
    ('templates', '模板表', [
        ('id', 'TEXT PK', '模板 ID'),
        ('name', 'TEXT NOT NULL', '模板名称'),
        ('description', 'TEXT', '模板描述'),
        ('preview_url', 'TEXT', '预览 URL'),
        ('is_active', 'INTEGER', '是否启用'),
        ('created_at', 'TEXT', '创建时间'),
        ('features', 'TEXT', '特性 JSON'),
    ]),
    ('audit_logs', '审计日志', [
        ('id', 'TEXT PK', '日志 ID'),
        ('action', 'TEXT NOT NULL', '操作类型'),
        ('target_type', 'TEXT NOT NULL', '目标类型'),
        ('target_id', 'TEXT', '目标 ID'),
        ('detail', 'TEXT', '详情'),
        ('ip', 'TEXT', '请求 IP'),
        ('created_at', 'TEXT', '创建时间'),
    ]),
]

for name, desc, columns in tables:
    doc.add_heading(f'{name} — {desc}', level=4)
    add_table(doc, ['列名', '类型', '说明'], [(c[0], c[1], c[2]) for c in columns])
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════
# 6. API DESIGN
# ══════════════════════════════
doc.add_heading('6. API 设计', level=1)
doc.add_paragraph('所有 API 使用 JSON 格式，除公开接口外需在 Header 中携带 `Authorization: Bearer {token}`。')

doc.add_heading('6.1 公开接口', level=2)
add_table(doc,
    ['方法', '路径', '说明', '速率限制'],
    [
        ['GET', '/api/health', '健康检查', '无'],
        ['POST', '/api/auth/login', '管理员登录', '5/min/IP'],
        ['POST', '/api/merchants/verify', '商户 Token 验证', '无'],
    ]
)

doc.add_heading('6.2 商户管理接口', level=2)
add_table(doc,
    ['方法', '路径', '说明', '速率限制'],
    [
        ['GET', '/api/merchants', '商户列表（支持 ?status= 筛选）', '无'],
        ['GET', '/api/merchants/:id', '商户详情', '无'],
        ['POST', '/api/merchants', '创建商户', '20/min/IP'],
        ['PUT', '/api/merchants/:id', '更新商户', '20/min/IP'],
        ['DELETE', '/api/merchants/:id', '删除商户（软删除）', '20/min/IP'],
        ['POST', '/api/merchants/:id/token', '重新生成 Token', '20/min/IP'],
    ]
)

doc.add_heading('6.3 部署接口', level=2)
add_table(doc,
    ['方法', '路径', '说明', '速率限制'],
    [
        ['GET', '/api/merchants/:id/deployments', '部署历史', '无'],
        ['POST', '/api/merchants/:id/deployments', '创建部署记录', '20/min/IP'],
        ['PUT', '/api/deployments/:id', '更新部署状态', '20/min/IP'],
        ['POST', '/api/merchants/:id/deploy-cf', 'CF 直连部署商户 Pages', '20/min/IP'],
    ]
)

doc.add_heading('6.4 模板接口', level=2)
add_table(doc,
    ['方法', '路径', '说明', '速率限制'],
    [
        ['GET', '/api/templates', '模板列表', '无'],
        ['POST', '/api/templates/scrape', '采集网站生成模板', '3/min/IP'],
        ['GET', '/api/templates/scrape/:jobId', '查询采集任务', '无'],
        ['GET', '/api/templates/scrape-jobs', '采集任务列表', '无'],
        ['GET', '/api/templates/scrape/:id/file?file=index.html', '查看模板文件', '无'],
    ]
)

doc.add_page_break()

# ══════════════════════════════
# 7. DEPLOYMENT ARCHITECTURE
# ══════════════════════════════
doc.add_heading('7. 部署架构', level=1)

doc.add_paragraph('域名配置：')
add_table(doc,
    ['域名', '目标', '说明'],
    [
        ['saas.roseai.ca', 'rose-saas-admin.pages.dev', 'Admin UI（反向代理）'],
        ['rose-saas-central-api.touchwant.workers.dev', 'Central API Worker', 'API 端点'],
        ['storefront-{id}.pages.dev', '商户 Storefront', '商户静态网站'],
        ['restaurant-api-{id}.{sub}.workers.dev', '商户 Worker', '商户 API 服务'],
    ]
)

doc.add_paragraph('CI/CD 流程：')
steps = [
    '开发者 push 代码到 GitHub main 分支',
    'GitHub Actions 自动触发部署工作流',
    'Job 1: wrangler deploy Central API Worker',
    'Job 2: wrangler pages deploy Admin UI',
    '部署成功完成后立即生效',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph('商户 CF 直连部署流程：')
steps = [
    '管理员在 Admin UI 为商户配置 CF Account ID + API Token',
    '点击 "CF 部署" 按钮',
    'Central API 读取商户信息和模板文件',
    '替换模板占位符为商户数据',
    '通过 Cloudflare Direct Upload API 部署到商户 Pages',
    '记录部署结果到 deployments 表',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_page_break()

# ══════════════════════════════
# 8. COST ESTIMATE
# ══════════════════════════════
doc.add_heading('8. 费用估算', level=1)

doc.add_heading('8.1 平台方成本（20 商户）', level=2)
add_table(doc,
    ['项目', '月费'],
    [
        ['Cloudflare Workers（中央后台）', '$0（免费额度内）'],
        ['Cloudflare D1（中央数据库）', '$0（免费额度内）'],
        ['Cloudflare Pages（Admin UI）', '$0'],
        ['域名 roseai.ca', '~$0.08'],
        ['总计', '~$0.08'],
    ]
)

doc.add_heading('8.2 每商户成本', level=2)
add_table(doc,
    ['项目', '月费'],
    [
        ['Cloudflare Workers', '$0（免费 10万请求/天）'],
        ['Cloudflare D1', '$0（免费 5GB）'],
        ['Cloudflare R2', '$0（免费 10GB）'],
        ['Cloudflare Vectorize', '$0（Beta 免费）'],
        ['Cloudflare Workers AI', '$0（免费 1万神经元/天）'],
        ['Cloudflare Pages', '$0'],
        ['Twilio 号码', '$1/月'],
        ['Twilio 通话（~100分钟）', '~$1.30'],
        ['小计', '~$2.30'],
    ]
)

doc.add_heading('8.3 总成本估算', level=2)
add_table(doc,
    ['商户数', '月成本', '主要支出'],
    [
        ['10 商户', '~$24', 'Twilio'],
        ['50 商户', '~$116', 'Twilio'],
        ['100 商户', '~$231', 'Twilio'],
        ['100 商户（仅文本客服）', '~$100', 'Twilio 号码月租'],
    ]
)

doc.add_page_break()

# ══════════════════════════════
# 9. FUTURE ROADMAP
# ══════════════════════════════
doc.add_heading('9. 未来规划', level=1)

add_table(doc,
    ['阶段', '功能'],
    [
        ['Phase 1（已完成）', '基础官网 + 模板系统 + 商户管理 + CF 直连部署'],
        ['Phase 2（进行中）', '在线下单 + 在线支付（微信/支付宝）+ 智能文本客服'],
        ['Phase 3', 'Google Drive 知识库 + RAG + 智能电话（Twilio）'],
        ['Phase 4', '商户自助后台 + 数据看板 + 营销工具（优惠券/满减）'],
        ['Phase 5', '多门店管理 + 聚合配送 + 供应链对接'],
        ['Phase 6', '模板市场（开发者上传模板 + 分成机制）'],
    ]
)

# ── Save ──
output_path = '/Users/sean/coding/roseai/rose-saas/Rose_SaaS_PRD.docx'
doc.save(output_path)
print(f'PRD generated: {output_path}')
